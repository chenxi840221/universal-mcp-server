"""Document reader module for handling PDF, Word, and Excel files."""

import os
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import PyPDF2
from docx import Document
import pandas as pd

logger = logging.getLogger(__name__)


class DocumentReader:
    """Handler for reading various document formats."""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._read_pdf,
            '.docx': self._read_docx,
            '.doc': self._read_docx,  # Will attempt to read .doc as .docx
            '.xlsx': self._read_excel,
            '.xls': self._read_excel,
            '.csv': self._read_csv
        }
    
    def read_document(self, file_path: str, format: str = "auto") -> Dict[str, Any]:
        """
        Read and extract content from a document.
        
        Args:
            file_path: Path to the document file
            format: Document format or 'auto' for auto-detection
        
        Returns:
            Dictionary containing extracted content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            raise ValueError(f"Path is not a file: {file_path}")
        
        # Auto-detect format if needed
        if format == "auto":
            format = file_path.suffix.lower()
        
        if format not in self.supported_formats:
            raise ValueError(f"Unsupported format: {format}. Supported formats: {list(self.supported_formats.keys())}")
        
        # Get file metadata
        stat = file_path.stat()
        metadata = {
            "file_path": str(file_path),
            "file_name": file_path.name,
            "file_size": stat.st_size,
            "format": format,
            "modified_time": stat.st_mtime,
            "created_time": stat.st_ctime
        }
        
        try:
            # Read the document content
            reader_func = self.supported_formats[format]
            content = reader_func(file_path)
            
            return {
                "success": True,
                "metadata": metadata,
                "content": content
            }
        
        except Exception as e:
            logger.error(f"Error reading document {file_path}: {e}")
            return {
                "success": False,
                "metadata": metadata,
                "error": str(e)
            }
    
    def _read_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Read PDF file content."""
        content = {
            "type": "pdf",
            "pages": [],
            "total_pages": 0,
            "text_content": ""
        }
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            content["total_pages"] = len(pdf_reader.pages)
            
            all_text = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    content["pages"].append({
                        "page_number": page_num + 1,
                        "text": page_text,
                        "character_count": len(page_text)
                    })
                    all_text.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    content["pages"].append({
                        "page_number": page_num + 1,
                        "text": "",
                        "error": str(e)
                    })
            
            content["text_content"] = "\n".join(all_text)
            content["total_characters"] = len(content["text_content"])
        
        return content
    
    def _read_docx(self, file_path: Path) -> Dict[str, Any]:
        """Read Word document content."""
        content = {
            "type": "docx",
            "paragraphs": [],
            "tables": [],
            "text_content": ""
        }
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            all_text = []
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    paragraph_data = {
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal"
                    }
                    content["paragraphs"].append(paragraph_data)
                    all_text.append(para.text)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = {
                    "table_number": table_idx + 1,
                    "rows": len(table.rows),
                    "columns": len(table.columns) if table.rows else 0,
                    "data": []
                }
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data["data"].append(row_data)
                
                content["tables"].append(table_data)
            
            content["text_content"] = "\n".join(all_text)
            content["total_paragraphs"] = len(content["paragraphs"])
            content["total_tables"] = len(content["tables"])
            content["total_characters"] = len(content["text_content"])
        
        except Exception as e:
            # If docx fails, it might be an older .doc format
            logger.warning(f"Failed to read as .docx: {e}")
            raise ValueError(f"Unable to read Word document. It might be in an older .doc format or corrupted: {e}")
        
        return content
    
    def _read_excel(self, file_path: Path) -> Dict[str, Any]:
        """Read Excel file content."""
        content = {
            "type": "excel",
            "sheets": [],
            "summary": {}
        }
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            content["summary"]["total_sheets"] = len(excel_file.sheet_names)
            content["summary"]["sheet_names"] = excel_file.sheet_names
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                sheet_data = {
                    "sheet_name": sheet_name,
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": list(df.columns),
                    "data": df.to_dict('records') if len(df) <= 1000 else [],  # Limit data for large files
                    "sample_data": df.head(10).to_dict('records') if len(df) > 1000 else [],
                    "data_types": df.dtypes.to_dict()
                }
                
                # Add summary statistics for numeric columns
                numeric_columns = df.select_dtypes(include=['number']).columns
                if len(numeric_columns) > 0:
                    sheet_data["numeric_summary"] = df[numeric_columns].describe().to_dict()
                
                content["sheets"].append(sheet_data)
            
            content["summary"]["total_rows"] = sum(sheet["rows"] for sheet in content["sheets"])
            content["summary"]["total_columns"] = sum(sheet["columns"] for sheet in content["sheets"])
        
        except Exception as e:
            logger.error(f"Error reading Excel file: {e}")
            raise ValueError(f"Unable to read Excel file: {e}")
        
        return content
    
    def _read_csv(self, file_path: Path) -> Dict[str, Any]:
        """Read CSV file content."""
        content = {
            "type": "csv",
            "data": [],
            "summary": {}
        }
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            df = None
            
            for encoding in encodings:
                try:
                    df = pd.read_csv(file_path, encoding=encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if df is None:
                raise ValueError("Unable to read CSV file with common encodings")
            
            content["summary"] = {
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": list(df.columns),
                "data_types": df.dtypes.to_dict()
            }
            
            # Include data (limit for large files)
            if len(df) <= 1000:
                content["data"] = df.to_dict('records')
            else:
                content["sample_data"] = df.head(100).to_dict('records')
                content["summary"]["note"] = "Large file - showing first 100 rows only"
            
            # Add summary statistics for numeric columns
            numeric_columns = df.select_dtypes(include=['number']).columns
            if len(numeric_columns) > 0:
                content["numeric_summary"] = df[numeric_columns].describe().to_dict()
        
        except Exception as e:
            logger.error(f"Error reading CSV file: {e}")
            raise ValueError(f"Unable to read CSV file: {e}")
        
        return content
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.supported_formats.keys())